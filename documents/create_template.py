"""
One-time script to generate the Word template (visa_letter.docx).

The template uses Jinja2 syntax that docxtpl understands:
  {{ variable }}   – simple substitution
  {% for ... %}    – loop (for crew rows)

Run once before starting the bot:
    python documents/create_template.py

After generation you can open the .docx file in Word/LibreOffice and
customise fonts, logos, headers, footers etc. — just keep the
{{ placeholders }} intact.
"""
import os
import sys

# Allow running directly from project root
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "templates", "visa_letter.docx"
)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background colour (e.g. '4472C4')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_white(paragraph) -> None:
    for run in paragraph.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def create_template() -> None:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── Header: date & reference ──────────────────────────────────────────────
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Date: {{ issue_date }}")
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # ── Recipient block ───────────────────────────────────────────────────────
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
    recipient.add_run(
        "To: The Director General of Immigration\n"
        "Port / Airport of {{ destination }}"
    ).font.size = Pt(11)

    doc.add_paragraph()

    # ── Subject line ──────────────────────────────────────────────────────────
    subject = doc.add_paragraph()
    subject.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subject.add_run("Subject: Request for {{ visa_type }} Visa – Vessel {{ ship_name }}")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ── Body ──────────────────────────────────────────────────────────────────
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(
        "Dear Sir/Madam,\n\n"
        "We respectfully request the issuance of a {{ visa_type }} visa for the crew members "
        "of the vessel detailed below, and kindly ask for your kind cooperation in this matter."
    ).font.size = Pt(11)

    doc.add_paragraph()

    # ── Vessel details table ──────────────────────────────────────────────────
    vessel_heading = doc.add_paragraph()
    run = vessel_heading.add_run("Vessel Details")
    run.bold = True
    run.font.size = Pt(11)

    vtbl = doc.add_table(rows=4, cols=2)
    vtbl.style = "Table Grid"
    vtbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    vessel_rows = [
        ("Ship Name", "{{ ship_name }}"),
        ("Ship Owner", "{{ ship_owner }}"),
        ("IMO Number", "{{ ship_imo }}"),
        ("Registration Date", "{{ ship_reg_date }}"),
    ]
    for i, (label, value) in enumerate(vessel_rows):
        vtbl.rows[i].cells[0].text = label
        vtbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        vtbl.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ── Routing details ───────────────────────────────────────────────────────
    route_heading = doc.add_paragraph()
    run = route_heading.add_run("Route Information")
    run.bold = True
    run.font.size = Pt(11)

    rtbl = doc.add_table(rows=2, cols=2)
    rtbl.style = "Table Grid"
    rtbl.rows[0].cells[0].text = "Origin"
    rtbl.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    rtbl.rows[0].cells[1].text = "{{ origin }}"
    rtbl.rows[1].cells[0].text = "Destination"
    rtbl.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    rtbl.rows[1].cells[1].text = "{{ destination }}"

    doc.add_paragraph()

    # ── Crew manifest ─────────────────────────────────────────────────────────
    crew_heading = doc.add_paragraph()
    run = crew_heading.add_run("Crew Manifest  (Total: {{ crew_count }} persons)")
    run.bold = True
    run.font.size = Pt(11)

    headers = ["#", "Full Name", "Passport No.", "Passport Expiry",
               "CDC No.", "CDC Expiry", "Gender", "Date of Birth", "Rank"]

    ctbl = doc.add_table(rows=1, cols=len(headers))
    ctbl.style = "Table Grid"

    # Header row
    hdr_row = ctbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        _set_cell_bg(cell, "4472C4")
        _bold_white(cell.paragraphs[0])
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    # Data rows — docxtpl Jinja2 loop block
    # We embed the loop markers as raw text in separate rows;
    # docxtpl will expand this at render time.
    for_row = ctbl.add_row()
    for_row.cells[0].text = "{%tr for member in crew %}"
    # Blank out remaining cells so docxtpl doesn't choke
    for i in range(1, len(headers)):
        for_row.cells[i].text = ""

    data_row = ctbl.add_row()
    data_row.cells[0].text = "{{ member.index }}"
    data_row.cells[1].text = "{{ member.full_name }}"
    data_row.cells[2].text = "{{ member.passport_number }}"
    data_row.cells[3].text = "{{ member.passport_expiry }}"
    data_row.cells[4].text = "{{ member.cdc_number }}"
    data_row.cells[5].text = "{{ member.cdc_expiry }}"
    data_row.cells[6].text = "{{ member.gender }}"
    data_row.cells[7].text = "{{ member.date_of_birth }}"
    data_row.cells[8].text = "{{ member.rank }}"

    endfor_row = ctbl.add_row()
    endfor_row.cells[0].text = "{%tr endfor %}"
    for i in range(1, len(headers)):
        endfor_row.cells[i].text = ""

    doc.add_paragraph()

    # ── Closing ───────────────────────────────────────────────────────────────
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    closing.add_run(
        "We confirm that the above-mentioned crew members are duly authorised personnel "
        "and request your kind assistance in processing their visas at the earliest convenience.\n\n"
        "Yours faithfully,"
    ).font.size = Pt(11)

    doc.add_paragraph("\n\n_________________________________")
    doc.add_paragraph("Authorised Agent / Shipping Company")

    # ── Save ──────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Template created: {OUTPUT_PATH}")


if __name__ == "__main__":
    create_template()
